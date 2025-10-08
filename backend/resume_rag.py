import os
import tempfile
from typing import List, Dict, Optional, cast
from dataclasses import dataclass
import logging
from pathlib import Path

# Document processing
import pdfplumber
from docx import Document
import markdown

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document as LangChainDocument
from langchain_google_genai import ChatGoogleGenerativeAI

# NLP processing
import spacy
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

logger = logging.getLogger(__name__)

@dataclass
class ResumeInfo:
    """Structured information extracted from resume"""
    candidate_id: str
    raw_text: str
    skills: List[str]
    experience: List[str]
    education: List[str]
    projects: List[str]
    summary: str
    years_experience: int = 0

class ResumeProcessor:
    """Handles resume file processing and text extraction"""
    
    def __init__(self):
        self._setup_nltk()
        self._setup_spacy()
    
    def _setup_nltk(self):
        """Download required NLTK data"""
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('tokenizers/punkt_tab')
            nltk.data.find('corpora/stopwords')
        except LookupError:
            try:
                nltk.download('punkt')
                nltk.download('punkt_tab')
                nltk.download('stopwords')
            except Exception as e:
                logger.warning(f"Failed to download NLTK data: {e}")
    
    def _setup_spacy(self):
        """Setup spaCy for NLP processing"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
            logger.info("Continuing without spaCy - using basic text processing")
            self.nlp = None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF resume"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX resume"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT resume"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            raise
    
    def process_resume_file(self, file_path: str, file_type: str) -> str:
        """Process resume file and extract text based on file type"""
        if file_type.lower() == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        elif file_type.lower() == 'txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

class ResumeAnalyzer:
    """Analyzes resume content and extracts structured information"""
    
    def __init__(self):
        self.processor = ResumeProcessor()
        self._load_skill_keywords()
        self._load_experience_keywords()
    
    def _load_skill_keywords(self):
        """Load common skill keywords for extraction"""
        self.skill_keywords = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask'],
            'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform'],
            'data_science': ['pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'spark'],
            'tools': ['git', 'jira', 'confluence', 'jenkins', 'gitlab', 'github'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'analytical']
        }
    
    def _load_experience_keywords(self):
        """Load keywords for experience extraction"""
        self.experience_keywords = [
            'worked', 'developed', 'designed', 'implemented', 'managed', 'led', 'created',
            'built', 'maintained', 'optimized', 'delivered', 'collaborated', 'architected'
        ]
    
    def extract_candidate_name(self, text: str) -> str:
        """Extract candidate name from resume text"""
        import re
        
        lines = text.split('\n')
        
        # Look for name patterns in first few lines
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if not line:
                continue
                
            # Skip common resume headers
            headers_to_skip = ['resume', 'curriculum vitae', 'cv', 'profile', 'contact', 'email', 'phone']
            if any(header in line.lower() for header in headers_to_skip):
                continue
                
            # Look for name patterns (2-4 words, mostly alphabetic)
            words = line.split()
            if 2 <= len(words) <= 4:
                # Check if it looks like a name (mostly alphabetic, proper case)
                if all(word.replace('-', '').replace("'", '').isalpha() and word[0].isupper() for word in words):
                    return line
        
        # Fallback: look for "Name:" pattern
        name_pattern = r'(?:name|full name)\s*:?\s*([A-Z][a-zA-Z\s\'-]+)'
        match = re.search(name_pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        
        return ""
    
    def extract_position_title(self, text: str) -> str:
        """Extract job position/title from resume text"""
        import re
        
        # Common job titles and patterns
        positions = [
            'software developer', 'software engineer', 'full stack developer', 'frontend developer', 
            'backend developer', 'web developer', 'mobile developer', 'devops engineer',
            'data scientist', 'data analyst', 'machine learning engineer', 'ai engineer',
            'product manager', 'project manager', 'business analyst', 'qa engineer',
            'system administrator', 'network engineer', 'security analyst', 'cloud engineer',
            'ui/ux designer', 'graphic designer', 'technical writer', 'scrum master'
        ]
        
        text_lower = text.lower()
        
        # Look for "seeking" or "looking for" patterns
        seeking_patterns = [
            r'seeking\s+(?:a\s+)?(?:position\s+as\s+)?([a-z\s]+?)(?:\.|\n|,)',
            r'looking\s+for\s+(?:a\s+)?(?:position\s+as\s+)?([a-z\s]+?)(?:\.|\n|,)',
            r'applying\s+for\s+(?:the\s+)?(?:position\s+of\s+)?([a-z\s]+?)(?:\.|\n|,)'
        ]
        
        for pattern in seeking_patterns:
            match = re.search(pattern, text_lower)
            if match:
                title = match.group(1).strip()
                # Check if it matches known positions
                for pos in positions:
                    if pos in title:
                        return pos.title()
        
        # Look for direct position mentions in first few lines
        lines = text.split('\n')[:10]
        for line in lines:
            line_lower = line.lower()
            for pos in positions:
                if pos in line_lower:
                    return pos.title()
        
        return "Software Developer"  # Default fallback
    
    def extract_experience_level(self, text: str) -> str:
        """Extract experience level from resume"""
        text_lower = text.lower()
        
        # Look for explicit experience mentions
        if any(word in text_lower for word in ['senior', '5+ years', '6+ years', '7+ years', 'lead', 'principal']):
            return "senior"
        elif any(word in text_lower for word in ['junior', '1 year', '2 years', 'entry level', 'fresh graduate']):
            return "entry"
        else:
            return "mid"  # Default to mid-level
    
    def extract_skills(self, text: str) -> List[str]:
        """Fast skill extraction with limited keywords"""
        text_lower = text.lower()
        found_skills = []
        
        # Only check most common skills for speed
        common_skills = [
            'python', 'java', 'javascript', 'react', 'node.js', 'sql', 'aws', 
            'docker', 'git', 'html', 'css', 'mongodb', 'postgresql', 'flask', 'django'
        ]
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return found_skills[:10]  # Limit to top 10
    
    def extract_experience(self, text: str) -> List[str]:
        """Extract work experience descriptions"""
        try:
            sentences = sent_tokenize(text)
        except LookupError:
            # Fallback: simple sentence splitting if NLTK punkt is not available
            sentences = text.split('. ')
        
        experience_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in self.experience_keywords):
                if len(sentence.split()) > 5:  # Only meaningful sentences
                    experience_sentences.append(sentence.strip())
        
        return experience_sentences[:10]  # Limit to top 10 experiences
    
    def extract_education(self, text: str) -> List[str]:
        """Extract education information"""
        education_keywords = ['university', 'college', 'degree', 'bachelor', 'master', 'phd', 'education']
        try:
            sentences = sent_tokenize(text)
        except LookupError:
            # Fallback: simple sentence splitting if NLTK punkt is not available
            sentences = text.split('. ')
        
        education_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in education_keywords):
                education_sentences.append(sentence.strip())
        
        return education_sentences[:5]
    
    def extract_projects(self, text: str) -> List[str]:
        """Extract project information"""
        project_keywords = ['project', 'developed', 'built', 'created', 'implemented']
        try:
            sentences = sent_tokenize(text)
        except LookupError:
            # Fallback: simple sentence splitting if NLTK punkt is not available
            sentences = text.split('. ')
        
        project_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in project_keywords):
                if len(sentence.split()) > 8:  # Only detailed project descriptions
                    project_sentences.append(sentence.strip())
        
        return project_sentences[:5]
    
    def estimate_experience_years(self, text: str) -> int:
        """Estimate years of experience from resume"""
        import re
        
        # Look for patterns like "X years", "X+ years", etc.
        patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'(\d+)\+?\s*yrs?\s*experience'
        ]
        
        text_lower = text.lower()
        max_years = 0
        
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                years = int(match)
                max_years = max(max_years, years)
        
        return max_years
    
    def analyze_resume(self, candidate_id: str, file_path: str, file_type: str) -> ResumeInfo:
        """Analyze resume and extract structured information"""
        try:
            # Extract text from file
            raw_text = self.processor.process_resume_file(file_path, file_type)
            
            # Extract structured information
            skills = self.extract_skills(raw_text)
            experience = self.extract_experience(raw_text)
            education = self.extract_education(raw_text)
            projects = self.extract_projects(raw_text)
            years_experience = self.estimate_experience_years(raw_text)
            
            # Generate summary
            summary = self._generate_summary(raw_text, skills, experience)
            
            return ResumeInfo(
                candidate_id=candidate_id,
                raw_text=raw_text,
                skills=skills,
                experience=experience,
                education=education,
                projects=projects,
                summary=summary,
                years_experience=years_experience
            )
            
        except Exception as e:
            logger.error(f"Error analyzing resume: {e}")
            raise
    
    def _generate_summary(self, text: str, skills: List[str], experience: List[str]) -> str:
        """Generate a summary of the resume"""
        summary_parts = []
        
        if skills:
            skills_str = ", ".join(skills[:5])  # Top 5 skills
            summary_parts.append(f"Skills: {skills_str}")
        
        if experience:
            summary_parts.append(f"Key experience: {experience[0][:100]}...")
        
        return ". ".join(summary_parts)

class ResumeRAG:
    """RAG system for resume-based question generation"""
    
    def __init__(self, google_api_key: str):
        self.google_api_key = google_api_key
        self.analyzer = ResumeAnalyzer()
        # Use lighter, faster embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",  # Faster, smaller model
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'batch_size': 1}  # Process one at a time for consistency
        )
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=google_api_key,
            temperature=0.3,
            max_tokens=100,
            timeout=10,        # Increased timeout to 10 seconds
            max_retries=2      # Limit retries to speed up failures
        )
        self.vector_store = None
        self.current_resume_info = None
    
    def process_resume_ultra_fast(self, candidate_id: str, storage_path: str, file_type: str) -> ResumeInfo:
        """Ultra-fast resume processing - no vector store, minimal processing"""
        try:
            from db_driver import DatabaseDriver
            db_driver = DatabaseDriver(use_service_role=True)
            
            # Download and process file quickly
            file_data = db_driver.download_resume(storage_path)
            
            import tempfile
            file_extension = file_type if file_type else storage_path.split('.')[-1]
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}')
            temp_file.write(file_data)
            temp_file.close()
            
            # Quick text extraction only
            raw_text = self.analyzer.processor.process_resume_file(temp_file.name, file_type)
            
            # Fast skill extraction with predefined list
            skills = self.analyzer.extract_skills(raw_text)
            
            # Create minimal resume info
            resume_info = ResumeInfo(
                candidate_id=candidate_id,
                raw_text=raw_text[:1000],  # Limit text for speed
                skills=skills,
                experience=["Software development experience"],  # Placeholder
                education=["Education background"],  # Placeholder
                projects=["Various projects"],  # Placeholder
                summary=f"Software developer with {len(skills)} technical skills",
                years_experience=2  # Default
            )
            
            self.current_resume_info = resume_info
            
            # Cleanup
            os.remove(temp_file.name)
            logger.info(f"Ultra-fast processing completed for: {candidate_id}")
            
            return resume_info
            
        except Exception as e:
            logger.error(f"Ultra-fast processing error: {e}")
            raise
    
    def process_resume(self, candidate_id: str, storage_path: str, file_type: str) -> ResumeInfo:
        """Alias for ultra-fast processing - maintains compatibility"""
        return self.process_resume_ultra_fast(candidate_id, storage_path, file_type)
    
    def _create_documents_from_resume(self, resume_info: ResumeInfo) -> List[LangChainDocument]:
        """Create LangChain documents from resume information"""
        documents = []
        
        # Add full text as main document
        documents.append(LangChainDocument(
            page_content=resume_info.raw_text,
            metadata={"type": "full_resume", "candidate_id": resume_info.candidate_id}
        ))
        
        # Add skills as separate documents
        if resume_info.skills:
            skills_text = f"Candidate skills: {', '.join(resume_info.skills)}"
            documents.append(LangChainDocument(
                page_content=skills_text,
                metadata={"type": "skills", "candidate_id": resume_info.candidate_id}
            ))
        
        # Add experience as separate documents
        for i, exp in enumerate(resume_info.experience):
            documents.append(LangChainDocument(
                page_content=f"Work experience: {exp}",
                metadata={"type": "experience", "candidate_id": resume_info.candidate_id, "exp_index": i}
            ))
        
        # Add projects as separate documents
        for i, project in enumerate(resume_info.projects):
            documents.append(LangChainDocument(
                page_content=f"Project: {project}",
                metadata={"type": "project", "candidate_id": resume_info.candidate_id, "project_index": i}
            ))
        
        return documents
    
    def generate_fast_questions(self, num_questions: int = 5) -> List[str]:
        """Generate questions instantly from resume without LLM - under 2 seconds"""
        if not self.current_resume_info:
            return [
                "Tell me about your software development experience.",
                "What programming languages are you most comfortable with?",
                "How do you approach problem-solving in code?",
                "Describe your experience with databases.",
                "What development tools do you use daily?"
            ][:num_questions]
        
        questions = []
        skills = self.current_resume_info.skills[:5]  # Top 5 skills
        experience = self.current_resume_info.experience[:3]  # Top 3 experiences
        
        # Generate skill-based questions instantly
        for skill in skills:
            if 'python' in skill.lower():
                questions.append(f"How many years of Python experience do you have?")
            elif 'javascript' in skill.lower() or 'js' in skill.lower():
                questions.append(f"Tell me about your JavaScript experience.")
            elif 'react' in skill.lower():
                questions.append(f"What React projects have you worked on?")
            elif 'java' in skill.lower() and 'javascript' not in skill.lower():
                questions.append(f"Describe your Java development experience.")
            elif 'sql' in skill.lower() or 'database' in skill.lower():
                questions.append(f"How do you handle database design and queries?")
            elif 'aws' in skill.lower() or 'cloud' in skill.lower():
                questions.append(f"What's your experience with cloud technologies?")
            else:
                questions.append(f"Tell me about your experience with {skill}.")
        
        # Add experience-based questions
        if experience:
            questions.append("Walk me through your most challenging project.")
            questions.append("How do you approach debugging complex issues?")
        
        # Add general software development questions
        questions.extend([
            "What's your preferred development methodology?",
            "How do you ensure code quality?",
            "Describe your experience with version control."
        ])
        
        # Return requested number of questions
        return questions[:num_questions]
    
    def _fallback_questions(self, position: str) -> List[str]:
        """Smart fallback questions based on resume if available"""
        if self.current_resume_info and self.current_resume_info.skills:
            # Create questions based on detected skills
            skills = self.current_resume_info.skills[:3]
            return [
                f"Tell me about your experience with {skills[0] if skills else 'your main technology'}.",
                f"How would you approach a challenging project using {skills[1] if len(skills) > 1 else 'your preferred tools'}?",
                "What's the most complex problem you've solved recently?",
                f"Describe a project where you used {skills[2] if len(skills) > 2 else 'your technical skills'}.",
                "What motivates you in your work?"
            ]
        else:
            # Generic fallback questions
            return [
                f"What interests you most about this {position} role?",
                "Can you walk me through your most challenging project?",
                "How do you stay updated with the latest technologies?",
                "Describe a time when you had to learn a new skill quickly.",
                "What are your career goals for the next 2-3 years?"
            ]
    
    def answer_question_with_context(self, question: str, user_response: str) -> str:
        """Provide follow-up based on resume context - simplified version"""
        if not self.current_resume_info:
            return "Please process a resume first."
        
        # Use basic resume context instead of vector search
        context = f"Skills: {', '.join(self.current_resume_info.skills[:3])}, Experience: {self.current_resume_info.years_experience} years"
        
        prompt = f"""
        Question: {question}
        Answer: {user_response}
        Candidate: {context}
        
        Give a brief follow-up question or comment (max 20 words).
        """
        
        try:
            response = self.llm.invoke(prompt)
            return cast(str, response.content)
        except Exception as e:
            logger.error(f"Error generating follow-up: {e}")
            logger.info("Using simple follow-up due to LLM timeout")
            # Smart fallback based on question type
            if "experience" in question.lower():
                return "Can you tell me more about the challenges you faced?"
            elif "project" in question.lower():
                return "What was your role in that project?"
            elif "skill" in question.lower() or "technology" in question.lower():
                return "How did you learn that technology?"
            else:
                return "Thank you. Let's continue with the next question."
    
    def cleanup(self):
        """Clean up vector store and temporary files"""
        if self.vector_store:
            try:
                # Clean up vector store directory
                import shutil
                if self.current_resume_info:
                    store_dir = f"./resume_vectorstore_{self.current_resume_info.candidate_id}"
                    if os.path.exists(store_dir):
                        shutil.rmtree(store_dir)
            except Exception as e:
                logger.error(f"Error cleaning up vector store: {e}")